#!/usr/bin/python

# USAGE:   $ python Nessus_algorithm.py <filename.nessus>

from sys import argv
import xml.etree.ElementTree as ET
import csv
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

#default file where excluded are stored
DEFAULT_EXCLUDED_FILENAME = "excludedIDs.csv"

class NessusFile:
  def __init__(self, tree):
    self._tree = tree

  def find_items(self):
    full_list = []

#finding our items in the given <filename>
    for host in self._tree.findall('Report/ReportHost'):
    #finding IP address
      self._ipaddr = host.find("HostProperties/tag/[@name='host-ip']").text

    #if existing - find CVE content
      for item in host.findall('ReportItem'):
        if item.find('cve') == None:
          cve = None
        else:
          cve = item.find('cve').text
        if item.find('see_also') == None:
          see_also = None
        else:
          see_also = item.find('see_also').text
    #other items
        risk_factor = item.find('risk_factor').text
        pluginID = item.get('pluginID')
        pluginName = item.get('pluginName')
        description = item.find('description').text
        solution = item.find('solution').text
    #applying priority to risk_factor by <order> and giving <full_list> ability to be sorted
        order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "None": 4}

        tuple_of_elements = (order[risk_factor], pluginName, risk_factor, cve, description, solution, see_also, pluginID)
        full_list.append(tuple_of_elements)

    full_list.sort()
    return full_list

#getter of IP address
  def get_ipaddr(self):
    return self._ipaddr

#getting excluded IDs from given or default file
class ExcludedLoader:
  def __init__(self, filename=None):
    #which file will be choosed condition
    if filename is not None:
      self._filename = filename
    else:
      self._filename = DEFAULT_EXCLUDED_FILENAME

  #getter of excluded ID dictionary
  def get_excluded(self):
    return self._get_excluded_ids_from_file(self._filename)

  #creating dictionary of IDs within given file
  def _get_excluded_ids_from_file(self, filename):
    with open(filename, 'r') as file:
      csv_file = csv.reader(file)
      for row in csv_file:
        dictionary = {i: True for i in row}
        return dictionary

#main processing class
class NessusProcess:
  def __init__(self, excluded_loader, ipaddr):
    self._excluded_loader = excluded_loader
    self._ipaddr = ipaddr

  #coloring risk factor cells in table function
  def _coloring_cells(self, cell, condition):
    if condition == 'Critical':
      color = '#960000'
    elif condition == 'High':
      color = '#fa3434'
    elif condition == 'Medium':
      color = '#ffa929'
    else:
      color = 'ffff40'
    tcPr = cell._tc.get_or_add_tcPr()
    tcVAlign = OxmlElement("w:shd")
    tcVAlign.set(qn("w:fill"), color)
    tcPr.append(tcVAlign)

#<name of variable=Document()>, <i> as second element in numeration, <title> as a title of paragraph
  #if '<i>' == 'NoneType' then we define APPENDIX paragraph but we still need to set a title
  def _create_main_paragraph(self, document, i, title):
    document.add_paragraph()
    paragraph = document.add_paragraph()
    if i != None:
      run = paragraph.add_run('1.' + str(i + 1) + ".     " + title)
      run.bold = True
      run.font.size = Pt(14)
    else:
      run = paragraph.add_run(title)
      run.bold = True
      run.font.size = Pt(16)
      run.font.color.rgb = RGBColor(255, 0 ,0)
    document.add_paragraph()
#<content> as position in a list, <number> - third element of numeration
  def _create_content_paragraph(self, document, i, content, number, title):
    paragraph = document.add_paragraph()
    run = paragraph.add_run('1.' + str(i + 1) + '.' + str(number) + ".   " + title)
    run.font.size = Pt(13)
    run.bold = True
    document.add_paragraph()
    paragraph = document.add_paragraph(content)
    paragraph.paragraph_format.left_indent = Cm(1.4)

    document.add_paragraph()

  #'.docx' creating function
  def process(self, full_list):
    excluded = self._excluded_loader.get_excluded() #excluded IDs
    ipaddr = self._ipaddr
    #list of elements defined in find_items() which priority is greater than 'None'
    result_list = [ x for x in full_list if x[1] not in excluded and x[0] < 4 ]

    #DOCUMENT TEMPLATE
    #it has header and footer
    path = 'document.docx'
    document = Document(path)
    for i in range(len(result_list)):
      self._create_main_paragraph(document, i, result_list[i][1])

      table = document.add_table(cols=2, rows=2)
      table.alignment = WD_TABLE_ALIGNMENT.RIGHT
      for row in table.rows:
        for cell in row.cells:
          cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      table.columns[0].width = Cm(7.0)
      table.columns[1].width = Cm(9.0)
      for row in table.rows:
        row.height = Cm(0.8)

      run = table.cell(0, 0).paragraphs[0].add_run('Rick Score:')
      run.bold = True
      run.font.size = Pt(12)
      table.cell(0, 1).paragraphs[0].add_run(result_list[i][2])
      self._coloring_cells(table.cell(0, 1), table.cell(0, 1).text)

      run10 = table.cell(1, 0).paragraphs[0].add_run('Affected Systems:')
      table.cell(1, 1).paragraphs[0].add_run('https://' + ipaddr)
      run10.bold = True
      run10.font.size = Pt(12)

      if result_list[i][3] != None:
        table.add_row()
        run = table.cell(2, 0).paragraphs[0].add_run('CVSS Risk Score:')
        table.cell(2, 1).paragraphs[0].add_run(result_list[i][3])
        run.bold = True
        run.font.size = Pt(12)
        table.rows[2].height = Cm(0.8)
        table.rows[2].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        table.rows[2].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
      document.add_paragraph()

      self._create_content_paragraph(document, i, result_list[i][4], 1, 'Vulnerability Description')

#TBA evidence generation function
      self._create_content_paragraph(document, i, '<TBA>', 2, 'Evidence')

      self._create_content_paragraph(document, i, result_list[i][5], 3, 'Recommendation')

      if result_list[i][7] != None:
        self._create_content_paragraph(document, i, result_list[i][6], 4, 'References')
      document.add_page_break()

    self._create_main_paragraph(document, None, '2.      APPENDIX')
    document.add_paragraph()
    paragraph21 = document.add_paragraph()
    run = paragraph21.add_run('2.1.    Port Scan Results')
    run.bold = True
    run.font.size = Pt(14)
    document.add_paragraph()
    table = document.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(8.0)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.cell(0, 0).text = ipaddr
#TBA open ports find function
    table.cell(0, 1).text = '<TBA>'
    table.rows[0].height = Cm(0.9)
    for row in table.rows:
      for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    document.save('Nessus-result.docx')

#execution function
def main(nessus_filename):
  """Main function"""
  tree = ET.parse(nessus_filename)

  nessus = NessusFile(tree)
  full_list = nessus.find_items()
  ipaddr = nessus.get_ipaddr()
  ld = ExcludedLoader()
  n = NessusProcess(ld, ipaddr)
  n.process(full_list)


if __name__ == '__main__':
  nessus_filename = argv[1]
  main(nessus_filename)